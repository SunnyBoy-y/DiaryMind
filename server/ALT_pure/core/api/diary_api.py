from fastapi import APIRouter, HTTPException, Body, Depends
from fastapi.responses import FileResponse
import os
from pathlib import Path
from pydantic import BaseModel
import docx
import json
from typing import List, Dict, Optional, Any
import re
import logging
from datetime import datetime
import sqlite3

from ..llm.qwen import LLM
from .auth_api import get_current_user
from .constants import (
    TIMELINE_PROMPT, PROBLEM_SOLUTION_PROMPT, KEY_POINTS_PROMPT, AUTO_STRUCTURE_PROMPT,
    EXTRACT_KEY_POINTS_PROMPT, ASSOCIATION_PROMPT, KNOWLEDGE_PROMPT, PLAN_PROMPT, REVIEW_PROMPT
)

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 初始化LLM实例
llm = LLM()

router = APIRouter()

# 获取项目根目录，更稳健的方式
try:
    # 尝试基于当前文件位置定位
    BASE_DIR = Path(__file__).resolve().parent.parent.parent.parent.parent
except Exception:
    # 回退到当前工作目录
    BASE_DIR = Path.cwd()

def get_user_storage_dir(user_id: int) -> Path:
    """
    获取用户的存储目录，如果不存在则创建。
    """
    storage_dir = BASE_DIR / "file_storage" / str(user_id)
    if not storage_dir.exists():
        storage_dir.mkdir(parents=True, exist_ok=True)
    return storage_dir

def _index_db_path() -> Path:
    p = BASE_DIR / "data"
    if not p.exists():
        p.mkdir(parents=True, exist_ok=True)
    return p / "diary_index.db"

def _ensure_index_schema(conn: sqlite3.Connection) -> None:
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS diaries (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id TEXT NOT NULL,
            filename TEXT NOT NULL,
            title TEXT,
            date TEXT,
            tags TEXT,
            mood TEXT,
            word_count INTEGER,
            preview TEXT,
            last_modified REAL,
            format TEXT
        )
        """
    )
    conn.execute(
        """
        CREATE UNIQUE INDEX IF NOT EXISTS idx_diaries_user_file
        ON diaries(user_id, filename)
        """
    )
    conn.commit()

def _safe_user_file_path(storage_dir: Path, filename: str) -> Path:
    """
    安全拼接并校验用户文件路径，防止路径穿越。
    """
    candidate = (storage_dir / filename).resolve()
    storage_dir_resolved = storage_dir.resolve()
    try:
        # Python 3.9+
        if not candidate.is_relative_to(storage_dir_resolved):
            raise HTTPException(status_code=400, detail="Invalid filename")
    except AttributeError:
        # 兼容早期版本
        if str(candidate).startswith(str(storage_dir_resolved)) is False:
            raise HTTPException(status_code=400, detail="Invalid filename")
    if candidate.is_dir():
        raise HTTPException(status_code=400, detail="Invalid filename")
    return candidate

class DiaryContent(BaseModel):
    filename: str
    content: str
    format: str = "md"  # md or docx

class StructureRequest(BaseModel):
    content: str
    template_type: str = "auto"  # auto, timeline, problem-solution, key-points

class KeyPointsRequest(BaseModel):
    content: str

class AssociationRequest(BaseModel):
    current_content: str
    history_content: str = ""

class LearningAssistantRequest(BaseModel):
    content: str
    type: str = "knowledge"  # knowledge, plan, review
    subject: str = "general"

class IntroRequest(BaseModel):
    title: Optional[str] = None
    tags: Optional[List[str]] = None
    mood: Optional[str] = None
    tone: Optional[str] = None
    hint: Optional[str] = None

class IntroResponse(BaseModel):
    intro: str

# --- Helper Functions ---

def _read_file_content(file_path: Path) -> tuple[str, str]:
    """
    读取文件内容，支持 .md, .txt, .docx。
    返回 (content, format)
    """
    ext = file_path.suffix.lower()
    content = ""
    
    if ext == ".docx":
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = "\n".join(full_text)
            return content, "docx"
        except Exception as e:
            logger.error(f"Error reading docx {file_path}: {e}")
            raise HTTPException(status_code=500, detail="Error reading docx")
    else:
        # Assume text/md
        try:
            content = file_path.read_text(encoding="utf-8")
            return content, ext.lstrip(".")
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            raise HTTPException(status_code=500, detail="Error reading file")

def _save_file_content(file_path: Path, content: str, fmt: str) -> None:
    """
    保存文件内容，支持 .md, .txt, .docx。
    """
    if fmt == "docx":
        try:
            doc = docx.Document()
            # Simple splitting by newlines for paragraphs
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(file_path)
        except Exception as e:
            logger.error(f"Error saving docx {file_path}: {e}")
            raise HTTPException(status_code=500, detail="Error saving docx")
    else:
        # Assume text/md
        try:
            file_path.write_text(content, encoding="utf-8")
        except Exception as e:
            logger.error(f"Error saving file {file_path}: {e}")
            raise HTTPException(status_code=500, detail="Error saving file")

def _get_diary_metadata(file_path: Path) -> Dict[str, Any]:
    """
    读取日记的元数据（Sidecar JSON）。
    """
    meta_path = file_path.with_name(f"{file_path.name}.meta.json")
    if meta_path.exists():
        try:
            return json.loads(meta_path.read_text(encoding="utf-8"))
        except Exception as e:
            logger.warning(f"Failed to read metadata for {file_path}: {e}")
    return {}

def _parse_diary_date(filename: str) -> str:
    """
    从文件名解析日期。
    格式：Diary_YYYY-MM-DD_HHMMSS.ext
    """
    date_str = ""
    if filename.startswith("Diary_"):
        date_part = filename[6:22]  # 提取 YYYY-MM-DD_HHMMSS 部分
        if len(date_part) >= 10:
            date_str = date_part[:10]  # 提取 YYYY-MM-DD 部分
    return date_str

def _extract_title(content: str, default: str) -> str:
    """
    从内容中提取标题（第一行 # 开头），否则返回默认值。
    """
    if content.strip():
        lines = content.strip().split("\n")
        for line in lines:
            if line.strip().startswith("# "):
                return line.strip()[2:]
    return default

# --- API Endpoints ---

@router.get("/list")
async def get_diary_list(current_user: dict = Depends(get_current_user)):
    """
    获取用户的日记列表。
    """
    storage_dir = get_user_storage_dir(current_user['id'])
    
    files = []
    for f in storage_dir.iterdir():
        if f.is_file() and f.suffix.lower() in {'.md', '.docx', '.txt'}:
            files.append(f.name)
    return files

@router.post("/generate-intro", response_model=IntroResponse)
async def generate_intro(request: IntroRequest, current_user: dict = Depends(get_current_user)):
    try:
        now = datetime.now().strftime("%Y-%m-%d")
        title = request.title or "新的开始"
        tags_text = ", ".join(request.tags) if request.tags else ""
        mood = request.mood or "平静"
        tone = request.tone or "温暖"
        hint = request.hint or ""
        prompt = (
            f"今天是{now}。请用{tone}的口吻，为题为《{title}》的日记生成一段不超过80字的个性化前言。"
            f"氛围倾向：{mood}。"
            f"标签：{tags_text}。"
            f"补充：{hint}。"
            f"要求：避免重复，用中文输出，适合开篇，简洁自然。"
        )
        text = await llm.chat(prompt, "你是贴心的写作助手")
        if not text:
            text = f"{now}，给自己一个轻松的开场。"
        clean = text.strip()
        if clean.startswith("```"):
            idx = clean.find("\n")
            if idx != -1:
                clean = clean[idx+1:]
        if clean.endswith("```"):
            clean = clean[:-3]
        clean = re.sub(r"<\s*script[\s\S]*?>[\s\S]*?<\s*/\s*script\s*>", "", clean, flags=re.IGNORECASE)
        clean = re.sub(r"[\r\n]+", " ", clean).strip()
        return IntroResponse(intro=clean)
    except Exception as e:
        logger.error(f"生成前言失败: {e}")
        raise HTTPException(status_code=500, detail="生成前言失败")

@router.get("/content/{filename}")
async def get_diary_content(filename: str, current_user: dict = Depends(get_current_user)):
    """
    获取指定日记的内容。
    """
    storage_dir = get_user_storage_dir(current_user['id'])
    file_path = _safe_user_file_path(storage_dir, filename)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    content, fmt = _read_file_content(file_path)
    return {"content": content, "format": fmt}

@router.post("/save")
async def save_diary(diary: DiaryContent, current_user: dict = Depends(get_current_user)):
    """
    保存日记内容。
    """
    storage_dir = get_user_storage_dir(current_user['id'])
        
    filename = diary.filename
    # Ensure extension matches format
    if not filename.lower().endswith(f".{diary.format}"):
        filename += f".{diary.format}"
        
    file_path = _safe_user_file_path(storage_dir, filename)
    _save_file_content(file_path, diary.content, diary.format)
    
    # 写入或更新元数据缓存，提升时光机性能
    meta_path = file_path.with_name(f"{file_path.name}.meta.json")
    preview = diary.content[:100] + ("..." if len(diary.content) > 100 else "")
    word_count = len(diary.content.replace("\n", " ").split())
    created_ts = int(Path(file_path).stat().st_mtime)
    created_date = datetime.fromtimestamp(created_ts).strftime("%Y-%m-%d")
    meta_data = {
        "summary": "",
        "tags": [],
        "mood": "",
        "mood_emoji": "",
        "word_count": word_count,
        "preview": preview,
        "created_at": created_ts,
        "created_at_date": created_date
    }
    try:
        meta_path.write_text(json.dumps(meta_data, ensure_ascii=False), encoding="utf-8")
    except Exception as e:
        logger.warning(f"Failed to write metadata for {file_path}: {e}")
            
    return {"status": "success", "filename": filename}

@router.post("/index/scan")
async def scan_index(current_user: dict = Depends(get_current_user)):
    try:
        storage_dir = get_user_storage_dir(current_user['id'])
        files = [f for f in storage_dir.iterdir() if f.is_file() and f.suffix.lower() in {'.md', '.docx', '.txt'}]
        conn = sqlite3.connect(str(_index_db_path()))
        _ensure_index_schema(conn)
        c = conn.cursor()
        count = 0
        for file in files:
            try:
                meta = _get_diary_metadata(file)
                if "preview" in meta and "word_count" in meta:
                    content_preview = meta.get("preview", "")
                    word_count = int(meta.get("word_count", 0))
                    ext = file.suffix.lstrip(".")
                else:
                    content, ext = _read_file_content(file)
                    content_preview = content[:100] + ("..." if len(content) > 100 else "")
                    word_count = len(content.replace("\n", " ").split())
                title = _extract_title(content_preview, file.name)
                date_str = _parse_diary_date(file.name) or meta.get("created_at_date") or datetime.fromtimestamp(file.stat().st_mtime).strftime("%Y-%m-%d")
                tags_text = json.dumps(meta.get("tags", []), ensure_ascii=False)
                mood = meta.get("mood", "")
                last_mod = file.stat().st_mtime
                c.execute(
                    """
                    INSERT INTO diaries(user_id, filename, title, date, tags, mood, word_count, preview, last_modified, format)
                    VALUES(?,?,?,?,?,?,?,?,?,?)
                    ON CONFLICT(user_id, filename) DO UPDATE SET
                        title=excluded.title,
                        date=excluded.date,
                        tags=excluded.tags,
                        mood=excluded.mood,
                        word_count=excluded.word_count,
                        preview=excluded.preview,
                        last_modified=excluded.last_modified,
                        format=excluded.format
                    """,
                    (str(current_user['id']), file.name, title, date_str, tags_text, mood, word_count, content_preview, last_mod, ext)
                )
                count += 1
            except Exception:
                continue
        conn.commit()
        conn.close()
        return {"indexed": count}
    except Exception as e:
        logger.error(f"索引扫描失败: {e}")
        raise HTTPException(status_code=500, detail="索引扫描失败")

@router.get("/meta/query")
async def query_meta(start: Optional[str] = None, end: Optional[str] = None, tag: Optional[str] = None, limit: int = 50, offset: int = 0, current_user: dict = Depends(get_current_user)):
    try:
        conn = sqlite3.connect(str(_index_db_path()))
        _ensure_index_schema(conn)
        c = conn.cursor()
        conditions = ["user_id = ?"]
        params: List[Any] = [str(current_user['id'])]
        if start:
            conditions.append("date >= ?")
            params.append(start)
        if end:
            conditions.append("date <= ?")
            params.append(end)
        if tag:
            conditions.append("tags LIKE ?")
            params.append(f"%{tag}%")
        where = " AND ".join(conditions)
        q = f"SELECT filename, title, date, tags, mood, word_count, preview, last_modified, format FROM diaries WHERE {where} ORDER BY date DESC, last_modified DESC LIMIT ? OFFSET ?"
        params.extend([limit, offset])
        rows = c.execute(q, params).fetchall()
        conn.close()
        result = []
        for r in rows:
            result.append({
                "filename": r[0],
                "title": r[1],
                "date": r[2],
                "tags": json.loads(r[3]) if r[3] else [],
                "mood": r[4],
                "word_count": r[5],
                "preview": r[6],
                "last_modified": r[7],
                "format": r[8]
            })
        return {"items": result, "count": len(result)}
    except Exception as e:
        logger.error(f"索引查询失败: {e}")
        raise HTTPException(status_code=500, detail="索引查询失败")

@router.delete("/delete/{filename}")
async def delete_diary(filename: str, current_user: dict = Depends(get_current_user)):
    """
    删除指定日记。
    """
    storage_dir = get_user_storage_dir(current_user['id'])
    file_path = _safe_user_file_path(storage_dir, filename)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    try:
        file_path.unlink()
        # 尝试删除元数据文件
        meta_path = file_path.with_name(f"{filename}.meta.json")
        if meta_path.exists():
            meta_path.unlink()
            
        return {"status": "success", "message": "File deleted successfully"}
    except Exception as e:
        logger.error(f"Error deleting file {filename}: {e}")
        raise HTTPException(status_code=500, detail="Error deleting file")

@router.post("/structure")
async def structure_diary_content(request: StructureRequest):
    """
    AI辅助的日记内容结构化处理接口
    """
    try:
        template_type = request.template_type
        content = request.content
        
        # 根据不同的模板类型选择prompt
        if template_type == "timeline":
            prompt = TIMELINE_PROMPT.format(content=content)
        elif template_type == "problem-solution":
            prompt = PROBLEM_SOLUTION_PROMPT.format(content=content)
        elif template_type == "key-points":
            prompt = KEY_POINTS_PROMPT.format(content=content)
        else:  # auto
            prompt = AUTO_STRUCTURE_PROMPT.format(content=content)
        
        # 调用LLM进行结构化处理
        structured_content = await llm.chat(prompt, "你是一位专业的内容结构化整理助手")
        
        if not structured_content:
            raise HTTPException(status_code=500, detail="LLM处理失败，返回空结果")
        
        return {"structured_content": structured_content, "template_type": template_type}
    except Exception as e:
        logger.error(f"结构化处理失败: {e}")
        raise HTTPException(status_code=500, detail="结构化处理失败")

@router.post("/extract-key-points")
async def extract_key_points(request: KeyPointsRequest):
    """
    AI驱动的日记要点智能挖掘系统接口
    """
    try:
        content = request.content
        prompt = EXTRACT_KEY_POINTS_PROMPT.format(content=content)
        
        # 调用LLM进行要点提取
        key_points = await llm.chat(prompt, "你是一位专业的信息提取分析师")
        
        if not key_points:
            raise HTTPException(status_code=500, detail="LLM处理失败，返回空结果")
        
        return {"key_points": key_points}
    except Exception as e:
        logger.error(f"要点提取失败: {e}")
        raise HTTPException(status_code=500, detail="要点提取失败")

@router.post("/associate")
async def get_associations(request: AssociationRequest):
    """
    AI智能联想功能接口，基于用户当前输入和历史内容提供相关主题的联想建议
    """
    try:
        prompt = ASSOCIATION_PROMPT.format(
            current_content=request.current_content,
            history_content=request.history_content
        )
        
        # 调用LLM生成联想建议
        associations = await llm.chat(prompt, "你是一位创意写作助手，擅长提供富有启发性的思路建议")
        
        if not associations:
            raise HTTPException(status_code=500, detail="LLM处理失败，返回空结果")
        
        return {"associations": associations}
    except Exception as e:
        logger.error(f"联想生成失败: {e}")
        raise HTTPException(status_code=500, detail="联想生成失败")

@router.get("/time-machine")
async def get_time_machine_data(current_user: dict = Depends(get_current_user)):
    """
    日记时光机功能接口，返回按时间线组织的日记数据
    """
    try:
        storage_dir = get_user_storage_dir(current_user['id'])
        if not storage_dir.exists():
            return {
                "timeline": [],
                "stats": {
                    "total_diaries": 0,
                    "recent_diaries": 0,
                    "average_length": 0
                }
            }
        
        # 获取所有日记文件
        files = []
        for f in storage_dir.iterdir():
            if f.is_file() and f.suffix.lower() in {'.md', '.docx', '.txt'}:
                files.append(f)
        
        # 按时间排序（文件名格式：Diary_YYYY-MM-DD_HHMMSS.ext）
        files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        
        # 组织成时间线数据
        timeline = []
        total_words = 0
        recent_diaries = 0
        recent_date = None
        
        for file in files:
            # 解析日期
            date_str = _parse_diary_date(file.name)
            
            # 获取内容
            try:
                meta_data = _get_diary_metadata(file)
                # 优先使用元数据缓存，避免每次读取全文
                if "preview" in meta_data and "word_count" in meta_data:
                    content = meta_data.get("preview", "")
                    ext = file.suffix.lstrip(".")
                    word_count = int(meta_data.get("word_count", 0))
                else:
                    content, ext = _read_file_content(file)
                    word_count = len(content.replace("\n", " ").split())
            except Exception:
                content = "无法读取文件内容"
                ext = file.suffix.lstrip(".")
                word_count = 0
            
            total_words += word_count
            
            # 统计最近7天的日记数量
            if recent_date is None:
                recent_date = file.stat().st_mtime
            else:
                days_diff = (recent_date - file.stat().st_mtime) / (60 * 60 * 24)
                if days_diff <= 7:
                    recent_diaries += 1
            
            # 提取标题
            title = _extract_title(content, file.name)
            
            # 读取元数据（如果前面未读取到）
            if 'meta_data' not in locals():
                meta_data = _get_diary_metadata(file)
            mood = None
            if "mood" in meta_data or "mood_emoji" in meta_data:
                mood = f"{meta_data.get('mood_emoji', '')} {meta_data.get('mood', '')}".strip()
            
            # 日期回退：优先元数据，其次文件时间
            if not date_str:
                date_str = meta_data.get("created_at_date") or datetime.fromtimestamp(file.stat().st_mtime).strftime("%Y-%m-%d")
            
            # 添加到时间线
            timeline.append({
                "filename": file.name,
                "date": date_str,
                "title": title,
                "mood": mood,
                "tags": meta_data.get("tags", []),
                "summary": meta_data.get("summary", ""),
                "word_count": word_count,
                "last_modified": file.stat().st_mtime,
                "format": ext,
                "preview": content[:100] + ("..." if len(content) > 100 else "")
            })
        
        # 计算统计信息
        total_diaries = len(timeline)
        average_length = round(total_words / total_diaries) if total_diaries > 0 else 0
        
        return {
            "timeline": timeline,
            "stats": {
                "total_diaries": total_diaries,
                "recent_diaries": recent_diaries,
                "average_length": average_length
            }
        }
    except Exception as e:
        logger.error(f"获取时光机数据失败: {e}")
        raise HTTPException(status_code=500, detail="获取时光机数据失败")

@router.post("/learn")
async def learning_assistant(request: LearningAssistantRequest):
    """
    AI学习助手功能接口，用于知识点整理、学习计划生成和复习提醒
    """
    try:
        content = request.content
        type_ = request.type  # use type_ to avoid shadowing built-in type
        subject = request.subject
        
        prompt = ""
        role = "你是一位专业的学习助手"
        
        if type_ == "knowledge":
            prompt = KNOWLEDGE_PROMPT.format(subject=subject, content=content)
            role = "你是一位专业的知识整理专家，擅长将复杂内容结构化"
        elif type_ == "plan":
            prompt = PLAN_PROMPT.format(subject=subject, content=content)
            role = "你是一位专业的学习规划师，擅长制定高效的学习计划"
        elif type_ == "review":
            prompt = REVIEW_PROMPT.format(subject=subject, content=content)
            role = "你是一位专业的记忆专家，擅长制定科学的复习计划"
        
        result = await llm.chat(prompt, role)
        
        if not result:
            raise HTTPException(status_code=500, detail="LLM处理失败，返回空结果")
        
        return {"result": result, "type": type_}
    except Exception as e:
        logger.error(f"学习助手处理失败: {e}")
        raise HTTPException(status_code=500, detail="学习助手处理失败")
